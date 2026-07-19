# -*- coding: utf-8 -*-
"""
Build comprehensive RTL Persian Word guide for Healan RAG chatbot.
Font: B Nazanin (falls back visually if missing on machine).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "rag-guide-assets"
OUT = ROOT / "راهنمای-ساخت-چت‌بات-RAG-Healan.docx"
FONT = "B Nazanin"
FONT_CODE = "Consolas"


def set_run_font(run, size=12, bold=False, color=None, code=False):
    name = FONT_CODE if code else FONT
    run.font.name = name
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_para(doc, text, size=12, bold=False, center=False, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, level=1):
    # Heading style + RTL
    p = doc.add_heading(level=level)
    set_rtl(p)
    # clear default runs
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13}
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p


def add_code(doc, code: str, caption: str | None = None):
    if caption:
        add_para(doc, caption, size=11, bold=True, color=RGBColor(0x47, 0x55, 0x69))
    p = doc.add_paragraph()
    # LTR for code readability
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "0")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    # light background via shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F5F9")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    run = p.add_run(code.strip("\n"))
    set_run_font(run, size=9, code=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_rtl(p)
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True)
        # header shade
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "DBEAFE")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            set_rtl(p)
            run = p.add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def add_image(doc, name, width_in=6.2, caption=None):
    path = ASSETS / name
    if not path.exists():
        add_para(doc, f"(تصویر موجود نیست: {name})", size=10, color=RGBColor(0xB9, 0x1C, 0x1C))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        add_para(doc, caption, size=10, center=True, color=RGBColor(0x47, 0x55, 0x69))


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    # document bidi
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    # set doc default RTL
    body = doc.element.body
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    # section may already have; append if missing
    if sectPr.find(qn("w:bidi")) is None:
        sectPr.append(bidi)


def build():
    doc = Document()
    configure_styles(doc)

    add_para(doc, "سامانه Healan — دستیار هوشمند مطب", size=22, bold=True, center=True,
             color=RGBColor(0x1F, 0x4E, 0x79))
    add_para(doc, "راهنمای کامل RAG + رزرو نوبت (تاریخ شمسی/نسبی و OTP)", size=15, bold=True, center=True)
    add_para(doc, "نسخه Markdown: Docs/RAG-Healan-Guide.md", size=11, center=True,
             color=RGBColor(0x47, 0x55, 0x69))
    doc.add_page_break()

    add_h(doc, "فهرست مطالب", 1)
    add_bullets(doc, [
        "فصل ۱ — هدف و رویکرد direct",
        "فصل ۲ — دو جریان: Ingest و Ask",
        "فصل ۳ — جداول SQL و فیلدها (نام دقیق)",
        "فصل ۴ — خواندن داده از بخش‌های مختلف دیتابیس",
        "فصل ۵ — ChromaDB (healan_rag)",
        "فصل ۶ — پایپ‌لاین Ingest + خلاصه‌ساز",
        "فصل ۷ — سرویس‌های Background (Python + C#)",
        "فصل ۸ — مثال کامل FAQ با کلاس‌ها",
        "فصل ۹ — مثال کامل بلاگ با کلاس‌ها",
        "فصل ۱۰ — سقف، OTP، لاگ",
        "فصل ۱۱ — مدل‌ها از UI و Ollama",
        "فصل ۱۲ — نقشه فایل‌ها و کلاس‌های C# / Python",
        "فصل ۱۳ — دیپلوی و عیب‌یابی",
        "فصل ۱۴ — رزرو نوبت از چت (جدا از RAG)",
        "فصل ۱۵ — احراز هویت رزرو (OTP موبایل)",
        "فصل ۱۶ — ذخیره OTP رزرو و عیب‌یابی",
    ])
    doc.add_page_break()

    add_h(doc, "فصل ۱ — هدف و رویکرد", 1)
    add_bullets(doc, [
        "پاسخ فقط از دانش تأییدشده در SQL.",
        "اگر شباهت کافی نبود، حدس نزند.",
        "مهمان سقف روزانه؛ بعد OTP.",
        "لاگ بدون کند کردن پاسخ (RabbitMQ).",
    ])
    add_para(doc, "حالت production: answer_mode=direct — جواب = metadata.answer نزدیک‌ترین سند در Chroma. LLM فقط برای خلاصه blog/review در Ingest.")
    add_code(doc, "سوال → Embedding → Chroma(healan_rag) → score≥آستانه → metadata.answer", "فلو مفهومی Ask:")

    add_h(doc, "فصل ۲ — دو جریان اصلی", 1)
    add_image(doc, "01-architecture.png", caption="شکل ۱ — معماری")
    add_para(doc, "جریان A — ایندکس (Background / استارت)", bold=True)
    add_code(doc, """SQL (۵ جدول)
  → HealanSqlDataSource.load()
  → enrich_documents_with_summaries(blog/review)
  → VectorStore.clear + upsert
  → UPDATE RagSettings.LastSyncedAt""")
    add_para(doc, "جریان B — پرسش (آنلاین)", bold=True)
    add_code(doc, """Portal Assistant
  → PortalPublicController.RagAsk
  → RagAskQueryHandler
  → RagQuotaCounter + RagQuotaHelper
  → RagPythonService.AskAsync
  → RagPipeline.ask → VectorStore.search → ask_direct
  → RagChatLogPublisher → Q_RagChatLog
  → RagChatLogConsumerService → RagChatLogs""")

    add_h(doc, "فصل ۳ — جداول SQL و فیلدها", 1)
    add_h(doc, "۳-۱. جدول RagKnowledgeItems", 2)
    add_para(doc, "Entity: RagKnowledgeItem · فایل: Healan.Domain/Portal/Entities/RagKnowledgeItem.cs · ToTable(\"RagKnowledgeItems\")", size=11)
    add_table(doc, ["فیلد", "نقش در RAG"], [
        ["RagKnowledgeItemId", "PK → id سند faq-{id}"],
        ["Question", "سوال اصلی"],
        ["QuestionSummary / Keywords / Topic / SimilarQuestions", "غنی‌سازی SearchText"],
        ["Answer", "جواب کاربر = metadata.answer"],
        ["SearchText", "متنی که Embed می‌شود = Document.content"],
        ["IsActive / IsDeleted", "فقط Active=1 و Deleted=0"],
        ["Priority / SortOrder", "اولویت و ترتیب UI"],
    ])
    add_h(doc, "۳-۲. جدول RagSettings (RagSettingId=1)", 2)
    add_para(doc, "Entity: RagSetting · Migration مدل‌ها: 20260718194500_AddRagModelSettings.cs", size=11)
    add_table(doc, ["فیلد", "پیش‌فرض", "نقش"], [
        ["RagSettingId", "1", "PK ثابت"],
        ["SyncIntervalMinutes", "10", "فاصله background sync پایتون"],
        ["SimilarityThresholdPercent", "55", "آستانه Ask (۰٫۵۵)"],
        ["PythonApiUrl", "http://python-rag:8000", "آدرس سرویس"],
        ["IsEnabled", "1", "خاموش/روشن ربات و sync"],
        ["GuestDailyLimit", "10", "سقف مهمان"],
        ["AuthenticatedDailyLimit", "200", "سقف لاگین"],
        ["EmbeddingModel", "heydariAI/persian-embeddings", "مدل بردار"],
        ["SummarizeModel", "qwen2.5:3b", "مدل خلاصه blog/review"],
        ["LastSyncedAt", "NULL", "آخرین ingest موفق"],
        ["UpdatedAt", "UtcNow", "زمان ویرایش"],
    ])
    add_h(doc, "۳-۳. جدول RagChatLogs", 2)
    add_para(doc, "Entity: RagChatLog", size=11)
    add_table(doc, ["فیلد", "نقش"], [
        ["RagChatLogId", "PK"],
        ["Question / Answer", "متن گفتگو"],
        ["MatchedKnowledgeItemId", "id سند مچ‌شده"],
        ["SimilarityScore", "نمره کسینوسی"],
        ["SourceType", "faq / setting / content / blog / review"],
        ["SessionId / GuestKey / IdentityUserId / PhoneNumber", "هویت و سهمیه"],
        ["WasAnswered / CreatedAt", "وضعیت و زمان"],
    ])
    add_h(doc, "۳-۴ تا ۳-۷. منابع محتوایی", 2)
    add_table(doc, ["جدول", "Entity", "فیلدهای کلیدی خوانده‌شده", "شرط"], [
        ["PortalSiteSettings", "PortalSiteSetting", "SettingKey, SettingValue, Description", "همه ردیف‌ها"],
        ["PortalContentItems", "PortalContentItem", "PortalContentItemId, SectionType, Title, Subtitle, Body", "IsPublished=1 AND IsDeleted=0"],
        ["BlogPosts", "BlogPost", "BlogPostId, Title, Excerpt, Body", "IsPublished=1 AND IsDeleted=0"],
        ["PatientReviews", "PatientReview", "PatientReviewId, DisplayName, ReviewText, Rating", "Status=2 AND IsDeleted=0"],
    ])

    add_h(doc, "فصل ۴ — خواندن داده از بخش‌های مختلف دیتابیس", 1)
    add_para(doc, "کلاس HealanSqlDataSource · فایل Python/app/data/healan_sql_source.py")
    add_para(doc, "یک اتصال SQL، پنج کوئری جدا، خروجی همه در یک list[Document] مخلوط — هنوز به Chroma نرفته.")
    add_code(doc, """def load(self) -> list[Document]:
    documents.extend(self._load_faq(...))
    documents.extend(self._load_settings(...))
    documents.extend(self._load_content(...))
    documents.extend(self._load_blogs(...))
    documents.extend(self._load_reviews(...))
    return documents""", "نقطه ورود:")
    add_table(doc, ["متد", "جدول", "id در Chroma", "source", "Embed", "پاسخ کاربر"], [
        ["_load_faq", "RagKnowledgeItems", "faq-{id}", "faq", "SearchText", "Answer"],
        ["_load_settings", "PortalSiteSettings", "setting-{key}", "setting", "Desc|Key|Value", "SettingValue"],
        ["_load_content", "PortalContentItems", "content-{id}", "content", "Title|Subtitle|Body", "ترکیب همان‌ها"],
        ["_load_blogs", "BlogPosts", "blog-{id}", "blog", "عنوان|خلاصه", "خلاصه"],
        ["_load_reviews", "PatientReviews", "review-{id}", "review", "نام|امتیاز|خلاصه", "نظر نام: خلاصه"],
    ])
    add_para(doc, "ساختار Document (Python/app/data/base.py): id + content (برای Embed) + metadata (حداقل answer, source, id). سیستم از قبل فیلتر source نمی‌زند؛ نزدیک‌ترین سند در یک کالکشن برنده است.")
    add_para(doc, "تنظیمات sync از همان SQL:", bold=True)
    add_bullets(doc, [
        "read_sync_settings() → SELECT TOP 1 * FROM RagSettings",
        "touch_last_synced() → UPDATE RagSettings SET LastSyncedAt = SYSUTCDATETIME() WHERE RagSettingId = 1",
    ])

    add_h(doc, "فصل ۵ — ChromaDB", 1)
    add_para(doc, "کلاس VectorStore · Python/app/rag/vector_store.py · Embedding: PersianEmbeddingFunction · embeddings.py")
    add_table(doc, ["تنظیم", "مقدار"], [
        ["Collection", "healan_rag"],
        ["Persist Docker", "/data/chroma"],
        ["Client", "chromadb.PersistentClient"],
        ["فاصله", "cosine"],
        ["Score در Ask", "1 - distance"],
    ])
    add_bullets(doc, [
        "هر Ingest: clear() (حذف کالکشن + ساخت دوباره) سپس upsert batch ۱۰۰تایی.",
        "Chroma همیشه عکس لحظه‌ای کامل دانش SQL است؛ آپدیت جزئی ردیف‌به‌ردیف ندارد.",
        "هر سند: ids=[doc.id], documents=[doc.content], metadatas=[doc.metadata].",
    ])
    add_image(doc, "02-ingest.png", caption="شکل ۲ — از SQL تا Chroma")

    add_h(doc, "فصل ۶ — پایپ‌لاین Ingest", 1)
    add_para(doc, "کلاس RagPipeline.ingest · pipeline.py")
    add_code(doc, """1) apply_rag_sql_overrides()          # EmbeddingModel / SummarizeModel از RagSettings
2) HealanSqlDataSource.load()         # ۵ منبع
3) enrich_documents_with_summaries()  # فقط blog و review
4) store.clear() + store.add_documents()
5) touch_last_synced()""")
    add_table(doc, ["فایل", "کلاس / تابع"], [
        ["pipeline.py", "RagPipeline.ingest / ask"],
        ["summarize_docs.py", "enrich_documents_with_summaries"],
        ["services/summarizer.py", "summarize_for_rag, fallback_summary, strip_html"],
        ["runtime_settings.py", "apply_rag_sql_overrides"],
        ["service.py", "init_rag, get_rag_pipeline, is_ingesting"],
    ])
    add_h(doc, "۶-۱. مثال عددی بلاگ فشار خون", 2)
    add_table(doc, ["فیلد BlogPosts", "مقدار"], [
        ["BlogPostId", "12"],
        ["Title", "علائم فشار خون بالا"],
        ["Body", "HTML بلند"],
        ["IsPublished", "1"],
    ])
    add_para(doc, "بعد از _load_blogs: id=blog-12 با raw_text کامل. بعد از summarize_for_rag: content=\"عنوان | خلاصه\" و metadata.answer=خلاصه. اگر Ollama قطع باشد fallback ~۶۰۰ کاراکتر.")

    add_h(doc, "فصل ۷ — سرویس‌های Background", 1)
    add_h(doc, "۷-۱. Python: run_background_sync", 2)
    add_para(doc, "فایل: Python/app/rag/background_sync.py · شروع از main.py lifespan با asyncio.create_task")
    add_bullets(doc, [
        "شرط: rag_sync_enabled + DATA_SOURCE=sqlserver + connection string.",
        "هر دور: صبر اگر is_ingesting → read_sync_settings از RagSettings.",
        "اگر IsEnabled=false → skip؛ وگرنه _sync_once.",
        "_sync_once: reset_rag_pipeline → apply_rag_sql_overrides → RagPipeline.ingest → touch_last_synced.",
        "خواب: SyncIntervalMinutes * 60 (پیش‌فرض ۱۰؛ هر دور از SQL override می‌شود).",
        "جداگانه در استارت: init_rag() یک بار ingest اولیه.",
    ])
    add_code(doc, """while not stop:
  wait if is_ingesting()
  sql = HealanSqlDataSource.read_sync_settings(...)
  if sql.is_enabled:
      _sync_once()   # full clear+upsert Chroma
  sleep(SyncIntervalMinutes * 60)""")
    add_h(doc, "۷-۲. C#: RagChatLogConsumerService", 2)
    add_para(doc, "کلاس: RagChatLogConsumerService : BackgroundService · فایل: Healan.Infrastructure/Portal/RagChatLogConsumerService.cs · DI: AddHostedService")
    add_bullets(doc, [
        "صف: QueueNames.RagChatLog = Q_RagChatLog",
        "پیام: RagChatLogMessage · ناشر: RagChatLogPublisher بعد از Ask",
        "کار: ساخت RagChatLog و INSERT در جدول RagChatLogs",
        "هدف: مسیر Ask معطل INSERT نشود",
    ])
    add_para(doc, "Ask آنلاین Background نیست. CRUD FAQ فقط SQL را عوض می‌کند؛ تا sync بعدی در Chroma دیده نمی‌شود.", size=11)

    add_h(doc, "فصل ۸ — مثال کامل FAQ: آدرس مطب", 1)
    add_image(doc, "03-ask-flow.png", caption="شکل ۳ — Ask")
    add_table(doc, ["فیلد RagKnowledgeItems", "مقدار نمونه"], [
        ["RagKnowledgeItemId", "4"],
        ["Question", "آدرس مطب کجاست؟"],
        ["Keywords", "آدرس، مکان، شوشتر"],
        ["Answer", "آدرس مطب: شوشتر، خیابان طالقانی، …"],
        ["SearchText", "ساخته با RagSearchTextBuilder.Build"],
        ["IsActive", "1"],
    ])
    add_para(doc, "بعد از ingest → Chroma id=faq-4 با metadata.answer = Answer")
    add_para(doc, "گام‌به‌گام Ask (کلاس‌ها به ترتیب):", bold=True)
    add_table(doc, ["#", "کلاس / متد", "کار"], [
        ["1", "Assistant/index.tsx", "ارسال سوال + guestKey"],
        ["2", "PortalPublicController.RagAsk", "POST PortalPublic/RagAsk"],
        ["3", "RagAskQueryHandler.Handle", "منطق Ask"],
        ["4", "خواندن RagSettings", "IsEnabled, limits, threshold, PythonApiUrl"],
        ["5", "RagQuotaCounter + RagQuotaHelper", "Redis rag:quota:g:{guest}:{day}"],
        ["6", "RagPythonService.AskAsync", "POST /api/v1/rag/ask answer_mode=direct"],
        ["7", "routers/rag.py → rag_ask", "FastAPI"],
        ["8", "RagPipeline.ask", "ارکستراسیون"],
        ["9", "VectorStore.search", "Embed سوال + cosine"],
        ["10", "ask_direct (direct_answer.py)", "برگرداندن metadata.answer"],
        ["11", "RagQuotaCounter.IncrementTodayAsync", "used++"],
        ["12", "RagChatLogPublisher.Publish", "Q_RagChatLog"],
        ["13", "RagChatLogConsumerService", "INSERT RagChatLogs SourceType=faq"],
    ])
    add_code(doc, """{
  "answer": "آدرس مطب: شوشتر، …",
  "wasAnswered": true,
  "similarityScore": 0.82,
  "matchedKnowledgeItemId": 4,
  "sourceType": "faq",
  "remainingCount": 7
}""", "پاسخ نمونه:")

    add_h(doc, "فصل ۹ — مثال کامل بلاگ: فشار خون", 1)
    add_para(doc, "Ask از نظر کلاس‌ها عین فصل ۸ است. تفاوت فقط محتوای ایندکس‌شده است.")
    add_table(doc, ["مرحله", "FAQ", "بلاگ"], [
        ["جدول SQL", "RagKnowledgeItems", "BlogPosts"],
        ["متد load", "_load_faq", "_load_blogs"],
        ["قبل از Embed", "بدون LLM", "enrich_documents_with_summaries"],
        ["id Chroma", "faq-4", "blog-12"],
        ["metadata.answer", "Answer ثابت", "خلاصه LLM"],
        ["sourceType", "faq", "blog"],
    ])
    add_bullets(doc, [
        "سوال: «فشار خون بالا چه علائمی دارد؟»",
        "جستجو روی کل healan_rag بدون فیلتر source",
        "اگر FAQ دقیقی نباشد blog-12 برنده می‌شود و خلاصه برمی‌گردد",
        "لاگ: SourceType=blog و MatchedKnowledgeItemId=12",
    ])

    add_h(doc, "فصل ۱۰ — سقف، OTP، لاگ", 1)
    add_image(doc, "04-quota-otp.png", caption="شکل ۴ — سقف و OTP")
    add_table(doc, ["کاربر", "فیلد", "کلید Redis"], [
        ["مهمان", "GuestDailyLimit=10", "rag:quota:g:{guestKey}:{yyyyMMdd}"],
        ["لاگین", "AuthenticatedDailyLimit=200", "rag:quota:u:{userId}:{day}"],
    ])
    add_bullets(doc, [
        "فایل‌ها: RagQuotaHelper.cs، RagQuotaCounter.cs — fallback COUNT روی RagChatLogs",
        "OTP سقف RAG: PortalOtpRequest / PortalOtpVerify · توکن portal_rag_access_token",
        "OTP رزرو نوبت جداست (فصل ۱۵): BookingOtpRequest / BookingOtpVerify · PortalAuthModal",
        "لاگ: RagChatLogPublisher → Q_RagChatLog → RagChatLogConsumerService → RagChatLogs",
    ])

    add_h(doc, "فصل ۱۱ — مدل‌ها از UI و Ollama", 1)
    add_para(doc, "UI: AssistantSettings.tsx · API: RagKnowledgeController.SettingGet/SettingSave · Handler: RagSettingGetQueryHandler / RagSettingSaveCommandHandler")
    add_table(doc, ["فیلد UI", "ستون SQL", "اثر بعد از sync"], [
        ["مدل Embedding", "EmbeddingModel", "clear+upsert با مدل جدید"],
        ["مدل خلاصه‌ساز", "SummarizeModel", "خلاصه blog/review دوباره"],
    ])
    add_code(doc, """sudo snap start ollama
ollama pull qwen2.5:3b
docker exec healan-python-rag curl -s http://host.docker.internal:11434/api/tags""")

    add_h(doc, "فصل ۱۲ — نقشه فایل‌ها و کلاس‌ها", 1)
    add_h(doc, "۱۲-۱. Python", 2)
    add_table(doc, ["مسیر", "کلاس / نماد"], [
        ["app/main.py", "lifespan → init_rag + run_background_sync"],
        ["app/data/healan_sql_source.py", "HealanSqlDataSource"],
        ["app/data/base.py", "Document, DataSource"],
        ["app/rag/pipeline.py", "RagPipeline"],
        ["app/rag/summarize_docs.py", "enrich_documents_with_summaries"],
        ["app/services/summarizer.py", "summarize_for_rag"],
        ["app/rag/runtime_settings.py", "apply_rag_sql_overrides"],
        ["app/rag/embeddings.py", "PersianEmbeddingFunction"],
        ["app/rag/vector_store.py", "VectorStore"],
        ["app/rag/direct_answer.py", "ask_direct"],
        ["app/rag/background_sync.py", "run_background_sync, _sync_once"],
        ["app/rag/service.py", "init_rag, get_rag_pipeline"],
        ["app/routers/rag.py", "rag_ask, rag_ingest, rag_status"],
    ])
    add_h(doc, "۱۲-۲. C#", 2)
    add_table(doc, ["کلاس", "نقش"], [
        ["RagKnowledgeItem / RagSetting / RagChatLog", "Entity جداول RAG"],
        ["BlogPost / PatientReview / PortalSiteSetting / PortalContentItem", "Entity منابع محتوا"],
        ["RagAskQueryHandler", "هسته Ask"],
        ["RagPythonService", "HTTP به Python"],
        ["RagQuotaHelper / RagQuotaCounter", "سقف"],
        ["RagChatLogPublisher / RagChatLogConsumerService", "لاگ async"],
        ["RagKnowledgeSeed / RagSearchTextBuilder", "seed و SearchText"],
        ["PortalPublicController / RagKnowledgeController", "endpointها"],
    ])
    add_h(doc, "۱۲-۳. Endpointها", 2)
    add_table(doc, ["متد", "مسیر"], [
        ["POST", "/Healan/api/v1/PortalPublic/RagAsk"],
        ["GET", "/Healan/api/v1/PortalPublic/RagQuota"],
        ["GET/POST", "/Healan/api/v1/RagKnowledge/SettingGet|SettingSave"],
        ["POST", "/Healan/api/v1/PortalPublic/BookingOtpRequest"],
        ["POST", "/Healan/api/v1/PortalPublic/BookingOtpVerify"],
        ["GET", "/Healan/api/v1/PortalPublic/BookingOpenSlots"],
        ["POST", "/Healan/api/v1/PortalPublic/BookingCreate"],
        ["POST", "{Python}/api/v1/rag/ask"],
        ["POST/GET", "{Python}/api/v1/rag/ingest|status"],
    ])

    add_h(doc, "فصل ۱۳ — دیپلوی و عیب‌یابی", 1)
    add_bullets(doc, [
        "stash appsettings → pull ai → apply-server-appsettings.sh",
        "build --no-cache: healan-webapi healan-clinic healan-portal python-rag",
        "up --force-recreate همان سرویس‌ها + nginx-gateway",
        "هرگز docker compose down کل استک و هرگز build --no-deps",
    ])
    add_table(doc, ["علامت", "احتمال"], [
        ["فقط عنوان بلاگ", "Ollama خاموش / summarize fail"],
        ["مدل عوض شد جواب قدیمی", "sync هنوز اجرا نشده"],
        ["FAQ جدید در چت نیست", "صبر SyncIntervalMinutes یا recreate python-rag"],
        ["کانتینر به Ollama نمی‌رسد", "host.docker.internal + snap start ollama"],
    ])

    add_h(doc, "فصل ۱۴ — رزرو نوبت از چت (جدا از RAG)", 1)
    add_para(doc, "اگر پیام درباره نوبت باشد، فرانت بدون RagAsk وارد رزرو می‌شود. با Ingest/Chroma اشتباه گرفته نشود.")
    add_code(doc, """پیام کاربر
  → parseBookingIntent (bookingIntent.ts)
  → kind=none → RagAsk
  → dateParseError → پیام خطا (بدون لاگین)
  → ensureBookingAuth → PortalAuthModal
  → BookingOpenSlots / Create / Cancel / Reschedule""")
    add_table(doc, ["فایل", "نقش"], [
        ["Assistant/index.tsx", "ارکستراسیون UI و دکمه‌های اسلات"],
        ["Assistant/bookingIntent.ts", "تشخیص kind + روز + ساعت (بدون LLM)"],
        ["Assistant/jalaliDate.ts", "تاریخ شمسی مطلق → dayKey میلادی"],
        ["PortalAuthModal.tsx", "OTP موبایل رزرو"],
        ["portalApi.ts", "فراخوانی PortalPublic رزرو"],
    ])
    add_h(doc, "۱۴-۱. انواع قصد", 2)
    add_table(doc, ["kind", "نمونه", "رفتار"], [
        ["list_slots", "نوبت‌های فردا / ۳۱ تیر", "لیست اسلات با دکمه"],
        ["book", "فردا ساعت ۱۸:۳۰ نوبت بگیر", "Create یا لیست همان روز"],
        ["my_bookings", "نوبت‌های من", "رزروهای فعال"],
        ["cancel", "لغو نوبت", "انتخاب برای لغو"],
        ["reschedule", "تغییر نوبت", "جابجایی اسلات"],
        ["none", "آدرس مطب؟", "می‌رود به RAG"],
    ])
    add_para(doc, "ضد مثبت کاذب: «لغو اشتراک» یا «جابجا کنید» بدون کلمه نوبت → none.", size=11)

    add_h(doc, "۱۴-۲. تاریخ شمسی مطلق", 2)
    add_table(doc, ["فرم", "مثال"], [
        ["Y/M/D", "1405/04/31 یا ۱۴۰۵-۴-۳۱"],
        ["D/M/Y", "31/04/1405"],
        ["روز + ماه + سال", "۳۱ تیر ۱۴۰۵"],
        ["روز + ماه بدون سال", "۱۵ تیر → سال جاری؛ اگر گذشته سال بعد"],
    ])
    add_bullets(doc, [
        "کتابخانه jalaali-js؛ خروجی dayKey میلادی YYYY-MM-DD برای API",
        "مرز حرفی برای تیر/مهر/دی تا «دیگر» به‌اشتباه دی نشود",
        "روز نامعتبر → dateParseError قبل از OTP",
    ])

    add_h(doc, "۱۴-۳. تاریخ نسبی و روز هفته", 2)
    add_table(doc, ["الگو", "مثال", "نتیجه"], [
        ["امروز / فردا / پس‌فردا", "نوبت‌های فردا", "offset 0 / 1 / 2"],
        ["N روز بعد", "۴ روز بعد، بعد از ۳ روز", "offset N (۰…۶۰)"],
        ["روز هفته + بعد", "شنبه بعد، بعد از یکشنبه", "Occurance بعدی (+۷ اگر امروز همان روز)"],
        ["این + روز هفته", "این شنبه", "نزدیک‌ترین همان روز"],
        ["فقط روز هفته", "نوبت‌های پنجشنبه؟", "مثل mode this؛ علائم ؟ نرمال می‌شوند"],
    ])
    add_para(doc, "اولویت: شمسی مطلق > نسبی. N>60 → خطا (نه برگشت بی‌صدا به امروز). «هفته بعد» هنوز پشتیبانی نمی‌شود.")

    add_h(doc, "۱۴-۴. مثال end-to-end", 2)
    add_bullets(doc, [
        "«نوبت‌های شنبه بعد» → list_slots → auth → OpenSlots → دکمه Create",
        "«۳۲ فروردین ۱۴۰۵ نوبت» → dateParseError بدون باز شدن مودال",
        "ساعت: ۱۸:۳۰ یا ساعت ۹ → timeHm؛ slotMatchesTime با Date محلی",
    ])

    add_h(doc, "فصل ۱۵ — احراز هویت رزرو (OTP موبایل)", 1)
    add_para(doc, "UI: PortalAuthModal · Backend: BookingOtpRequest / BookingOtpVerify / CompleteProfile")
    add_code(doc, "phone → otp → confirm | profile → ادامه pendingIntent")
    add_table(doc, ["مرحله", "کار", "API"], [
        ["phone", "موبایل 09…", "BookingOtpRequest"],
        ["otp", "کد پیامک", "BookingOtpVerify → JWT + بیمار"],
        ["confirm", "بیمار کامل → تأیید", "ادامه رزرو"],
        ["profile", "نام/نام‌خانوادگی/کدملی", "BookingCompleteProfile"],
    ])
    add_table(doc, ["مقدار", "عدد"], [
        ["TTL کد OTP", "۵ دقیقه"],
        ["Cooldown ارسال مجدد", "۶۰ ثانیه"],
        ["Session بعد از verify", "۳۰ دقیقه"],
    ])
    add_para(doc, "اگر SMS شکست بخورد: حذف OTP + ClearCooldownAsync تا کاربر دوباره بتواند درخواست بدهد.")
    add_para(doc, "تفاوت با OTP سقف RAG (فصل ۱۰): توکن portal_rag_access_token فقط برای Ask است؛ JWT رزرو برای Create/Cancel است.", size=11)

    add_h(doc, "فصل ۱۶ — ذخیره OTP رزرو و عیب‌یابی", 1)
    add_para(doc, "قرارداد: IBookingOtpStore · حافظه: MemoryBookingOtpStore · SQL: DbBookingOtpStore روی جدول BookingAuthTokens")
    add_code(doc, """SetAsync / TryGetAsync / RemoveAsync
SetCooldownAsync / IsInCooldownAsync / ClearCooldownAsync
SetSessionAsync / GetSessionPhoneAsync""")
    add_bullets(doc, [
        "DbBookingOtpStore: اگر SQL/جدول در دسترس نباشد → fallback حافظه",
        "Get: اول SQL، بعد memory (برای writeهایی که قبلاً fallback شده‌اند)",
        "RedisBookingOtpStore هم الگوی مشابه دارد؛ DI فعال معمولاً DB است",
    ])
    add_table(doc, ["علامت", "اقدام"], [
        ["ارسال کد ناموفق + cooldown گیر", "recreate healan-webapi (ClearCooldown)"],
        ["502 روی BookingOtpRequest", "apply-server-appsettings + recreate API"],
        ["Verify بین instance کد را نمی‌بیند", "migrate جدول BookingAuthTokens"],
        ["تاریخ شمسی کار نمی‌کند", "--no-cache healan-portal + nginx؛ مارکر build-v13"],
    ])

    add_h(doc, "جمع‌بندی", 1)
    add_bullets(doc, [
        "دانش در جداول SQL با فیلدهای مشخص",
        "HealanSqlDataSource.load → پنج منبع → Document",
        "خلاصه فقط blog/review",
        "VectorStore روی healan_rag",
        "Background Python برای sync؛ Background C# برای لاگ",
        "Ask = Embedding + cosine + metadata.answer",
        "رزرو نوبت جدا: intent فارسی + شمسی/نسبی + OTP موبایل (DB/Memory)",
    ])

    doc.save(str(OUT))
    print("OK:", OUT)


if __name__ == "__main__":
    build()
