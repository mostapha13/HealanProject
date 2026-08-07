from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.responses import Response
from io import BytesIO
from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from pydantic import BaseModel, Field
from typing import Any
import re
import chromadb
import hashlib
import os
import difflib
import logging
import json

app = FastAPI(title="NegareshAI", version="0.1.0")
class HashEmbedding:
    def __call__(self, input):
        return [self._vector(str(value)) for value in input]
    @staticmethod
    def _vector(value: str) -> list[float]:
        raw = hashlib.sha256(value.encode('utf-8')).digest()
        return [((raw[i % len(raw)] / 255.0) * 2.0) - 1.0 for i in range(128)]

class SemanticEmbedding:
    def __init__(self, model_id: str):
        from sentence_transformers import SentenceTransformer
        self.model_id = model_id
        model_source = local_model_snapshot(model_id)
        self.model = SentenceTransformer(
            model_source,
            cache_folder=os.getenv("SENTENCE_TRANSFORMERS_HOME", "/models"),
            local_files_only=os.getenv("MODEL_OFFLINE", "true").lower() == "true")
    def __call__(self, input):
        values = [str(value) for value in input]
        return self.model.encode(
            values, normalize_embeddings=True, show_progress_bar=False).tolist()

def local_model_snapshot(model_id: str) -> str:
    """Resolve a cached Hugging Face model to a concrete local snapshot.

    Passing the repository id to recent transformers versions can still trigger
    a metadata request even with local_files_only enabled. A concrete snapshot
    path keeps the contract workflow fully offline once the model is cached.
    """
    if os.path.isdir(model_id):
        return model_id
    cache_root = os.getenv("SENTENCE_TRANSFORMERS_HOME", "/models")
    repository = os.path.join(cache_root, f"models--{model_id.replace('/', '--')}")
    main_ref = os.path.join(repository, "refs", "main")
    try:
        with open(main_ref, encoding="utf-8") as ref:
            revision = ref.read().strip()
        snapshot = os.path.join(repository, "snapshots", revision)
        if revision and os.path.isdir(snapshot):
            return snapshot
    except OSError:
        pass
    return model_id

_embeddings: dict[str, Any] = {}
def embedding_for(model_id: str):
    backend = os.getenv("EMBEDDING_BACKEND", "semantic")
    cache_key = f"{backend}:{model_id}"
    if cache_key not in _embeddings:
        _embeddings[cache_key] = (HashEmbedding() if backend == "hash"
                                  else SemanticEmbedding(model_id))
    return _embeddings[cache_key]

vector_client = chromadb.PersistentClient(path=os.getenv('CHROMA_PERSIST_DIR', '/data/chroma'))
RAG_COLLECTION = os.getenv("RAG_COLLECTION", "negareshai_documents")
DEFAULT_EMBEDDING_MODEL = "BAAI/bge-m3"
logger = logging.getLogger("negareshai-ai")

class RagContext(BaseModel):
    organizationId: str = Field(min_length=36, max_length=36)
    documentId: str = Field(min_length=36, max_length=36)
    versionId: str = Field(min_length=36, max_length=36)
    embeddingModel: str = Field(default=DEFAULT_EMBEDDING_MODEL, min_length=3)
    accessScope: str = Field(default="restricted", pattern="^(organization|restricted)$")
    allowedUserIds: list[str] = Field(default_factory=list)
    allowedGroupIds: list[str] = Field(default_factory=list)
    approvalState: str = Field(default="final", pattern="^final$")

class PipelineRequest(RagContext):
    fileName: str
    contentBase64: str
    maxChars: int = Field(default=1800, ge=200, le=10000)
    publishToRag: bool = False

class IndexRequest(RagContext):
    chunks: list[dict[str, Any]] = Field(min_length=1)

class DeleteVersionRequest(BaseModel):
    organizationId: str = Field(min_length=36, max_length=36)
    documentId: str = Field(min_length=36, max_length=36)
    versionId: str = Field(min_length=36, max_length=36)
    embeddingModel: str = Field(default=DEFAULT_EMBEDDING_MODEL, min_length=3)

class SearchRequest(BaseModel):
    organizationId: str = Field(min_length=36, max_length=36)
    userId: str = Field(min_length=1, max_length=200)
    groupIds: list[str] = Field(default_factory=list)
    query: str = Field(min_length=1)
    documentIds: list[str] | None = None
    limit: int = Field(default=5, ge=1, le=20)
    embeddingModel: str = Field(default=DEFAULT_EMBEDDING_MODEL, min_length=3)

@app.get("/health")
def health():
    return {"service": "negareshai-ai", "status": "healthy"}

@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    data = await file.read()
    name = (file.filename or "").lower()
    try:
        if name.endswith(".pdf"):
            reader = PdfReader(BytesIO(data))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif name.endswith(".docx"):
            document = Document(BytesIO(data))
            text = "\n".join(p.text for p in document.paragraphs)
        else:
            raise HTTPException(415, "Only PDF and DOCX are supported")
        return {"fileName": file.filename, "text": text, "characters": len(text), "ocrRequired": not bool(text.strip())}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, f"Document extraction failed: {exc}") from exc

def ocr_pdf_page(data: bytes, page_index: int) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image
        document = fitz.open(stream=data, filetype="pdf")
        page = document.load_page(page_index)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
        image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
        languages = os.getenv("OCR_LANGUAGES", "fas+eng")
        return pytesseract.image_to_string(image, lang=languages).strip()
    except Exception as exc:
        raise RuntimeError(f"OCR failed for page {page_index + 1}: {exc}") from exc

def extract_pages(data: bytes, name: str, enable_ocr: bool = True) -> list[dict]:
    if name.lower().endswith(".pdf"):
        reader = PdfReader(BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            used_ocr = False
            if enable_ocr and not text:
                text = ocr_pdf_page(data, index)
                used_ocr = bool(text)
            pages.append({"page": index + 1, "text": text, "ocr": used_ocr})
        return pages
    if name.lower().endswith(".docx"):
        document = Document(BytesIO(data))
        return [{"page": 1, "text": "\n".join(p.text for p in document.paragraphs)}]
    if name.lower().endswith((".jpg", ".jpeg", ".png", ".tif", ".tiff")):
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(BytesIO(data))
            text = pytesseract.image_to_string(
                image, lang=os.getenv("OCR_LANGUAGES", "fas+eng")).strip()
            return [{"page": 1, "text": text, "ocr": True}]
        except Exception as exc:
            raise HTTPException(422, f"Image OCR failed: {exc}") from exc
    raise HTTPException(415, "Only PDF, DOCX, JPG, PNG and TIFF are supported")

def structural_chunks(text: str, max_chars: int = 1800, page: int = 1) -> list[dict]:
    sections = re.split(r"(?m)(?=^\s*(?:ماده|بند|فصل|تبصره|[0-9۰-۹]+[.)-])\s*)", text)
    chunks: list[dict] = []
    for section in sections:
        section = section.strip()
        if not section:
            continue
        for offset in range(0, len(section), max_chars):
            value = section[offset:offset + max_chars].strip()
            if value:
                chunks.append({"text": value, "index": len(chunks), "section": section[:120], "page": page})
    return chunks

def collection_for(model_id: str):
    backend = os.getenv("EMBEDDING_BACKEND", "semantic")
    collection_key = f"{backend}:{model_id}"
    suffix = hashlib.sha256(collection_key.encode("utf-8")).hexdigest()[:12]
    return vector_client.get_or_create_collection(
        f"{RAG_COLLECTION}_{suffix}",
        embedding_function=embedding_for(model_id),
        metadata={"embeddingModel": model_id, "embeddingBackend": backend})

def normalize_persian(value: str) -> str:
    table = str.maketrans({
        **{character: str(index) for index, character in enumerate("۰۱۲۳۴۵۶۷۸۹")},
        "ي": "ی",
        "ك": "ک",
    })
    return re.sub(r"\s+", " ", value.translate(table).replace("٬", "").replace(",", "")).strip().lower()

def numeric_tokens(value: str) -> set[str]:
    return set(re.findall(r"\d+(?:[./-]\d+)*", normalize_persian(value)))

def lexical_score(query: str, document: str) -> float:
    query_tokens = set(re.findall(r"[\w./-]+", normalize_persian(query)))
    document_tokens = set(re.findall(r"[\w./-]+", normalize_persian(document)))
    return len(query_tokens & document_tokens) / max(len(query_tokens), 1)

def index_document_chunks(context: RagContext, chunks: list[dict]) -> int:
    if not chunks:
        return 0
    collection = collection_for(context.embeddingModel)
    prefix = f"{context.organizationId}:{context.documentId}:{context.versionId}:"
    existing = collection.get(where={
        "$and": [
            {"organizationId": context.organizationId},
            {"documentId": context.documentId},
            {"versionId": context.versionId}
        ]
    })
    if existing.get("ids"):
        collection.delete(ids=existing["ids"])
    ids = [f"{prefix}{index}" for index in range(len(chunks))]
    documents = [str(chunk["text"]) for chunk in chunks]
    metadata = [{
        "organizationId": context.organizationId,
        "documentId": context.documentId,
        "versionId": context.versionId,
        "page": int(chunk.get("page", 1)),
        "chunkIndex": index,
        "section": str(chunk.get("section", ""))[:500],
        "numbers": "|".join(sorted(numeric_tokens(str(chunk["text"]))))[:1000],
        "accessScope": context.accessScope,
        "allowedUserIds": "|".join(sorted(set(context.allowedUserIds)))[:4000],
        "allowedGroupIds": "|".join(sorted(set(context.allowedGroupIds)))[:4000],
        "embeddingModel": context.embeddingModel
        , "approvalState": context.approvalState
    } for index, chunk in enumerate(chunks)]
    collection.upsert(ids=ids, documents=documents, metadatas=metadata)
    return len(chunks)

@app.post("/chunk")
async def chunk_document(payload: dict):
    text = payload.get("text", "")
    if not isinstance(text, str) or not text.strip():
        raise HTTPException(400, "text is required")
    max_chars = int(payload.get("maxChars", 1800))
    if max_chars < 200 or max_chars > 10000:
        raise HTTPException(400, "maxChars must be between 200 and 10000")
    chunks = structural_chunks(text, max_chars)
    return {"chunks": chunks, "count": len(chunks)}

@app.post("/rag/index")
async def index_chunks(payload: IndexRequest):
    return {"collection": collection_for(payload.embeddingModel).name,
            "indexed": index_document_chunks(payload, payload.chunks)}

@app.post("/rag/delete-version")
async def delete_version(payload: DeleteVersionRequest):
    collection = collection_for(payload.embeddingModel)
    existing = collection.get(where={"$and": [
        {"organizationId": payload.organizationId},
        {"documentId": payload.documentId},
        {"versionId": payload.versionId}
    ]})
    ids = existing.get("ids") or []
    if ids:
        collection.delete(ids=ids)
    return {"deleted": len(ids)}

@app.post("/rag/search")
async def search_chunks(payload: SearchRequest):
    collection = collection_for(payload.embeddingModel)
    filters: list[dict[str, Any]] = [
        {"organizationId": payload.organizationId}, {"approvalState": "final"}]
    if payload.documentIds:
        filters.append({"documentId": {"$in": payload.documentIds}})
    where = filters[0] if len(filters) == 1 else {"$and": filters}
    candidate_count = collection.count()
    if candidate_count == 0:
        return {"results": [], "embeddingModel": payload.embeddingModel}
    result = collection.query(
        query_texts=[payload.query.strip()], n_results=candidate_count,
        where=where, include=["documents", "metadatas", "distances"])
    documents = result.get("documents", [[]])[0]
    metadatas = result.get("metadatas", [[]])[0]
    distances = result.get("distances", [[]])[0]
    query_numbers = numeric_tokens(payload.query)
    caller_groups = set(payload.groupIds)
    permitted = []
    for text, metadata, distance in zip(documents, metadatas, distances):
        allowed_users = set(filter(None, str(metadata.get("allowedUserIds", "")).split("|")))
        allowed_groups = set(filter(None, str(metadata.get("allowedGroupIds", "")).split("|")))
        if (metadata.get("accessScope") == "organization"
                or payload.userId in allowed_users
                or bool(caller_groups & allowed_groups)):
            permitted.append((text, metadata, distance))
    ranked = [{
        "text": text,
        "distance": distance,
        "score": (1.0 - float(distance))
                 + (0.25 * lexical_score(payload.query, text))
                 + (0.5 if query_numbers and query_numbers.issubset(numeric_tokens(text)) else 0.0),
        "citation": {
            "documentId": metadata["documentId"],
            "versionId": metadata["versionId"],
            "page": metadata["page"],
            "section": metadata.get("section", "")
        }
    } for text, metadata, distance in permitted]
    ranked.sort(key=lambda item: item["score"], reverse=True)
    return {"results": ranked[:payload.limit], "embeddingModel": payload.embeddingModel}

@app.post("/pipeline/process")
async def process_document(payload: PipelineRequest):
    import base64
    try:
        data = base64.b64decode(payload.contentBase64, validate=True)
        pages = extract_pages(data, payload.fileName)
        chunks: list[dict] = []
        for page in pages:
            page_chunks = structural_chunks(page["text"], payload.maxChars, page["page"])
            for chunk in page_chunks:
                chunk["index"] = len(chunks)
                chunks.append(chunk)
        if not chunks:
            return {"status": "ocr_required", "pageCount": len(pages),
                    "characters": 0, "chunkCount": 0}
        indexed = index_document_chunks(payload, chunks) if payload.publishToRag else 0
        return {"status": "ready" if payload.publishToRag else "extracted", "pageCount": len(pages),
                "characters": sum(len(page["text"]) for page in pages),
                "chunkCount": indexed,
                "ocrPageCount": sum(1 for page in pages if page.get("ocr")),
                "extractedText": "\f".join(page["text"] for page in pages)}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, f"Document processing failed: {exc}") from exc

def _currency_label(value: object) -> str:
    return {"IRR": "ریال", "IRT": "تومان", "USD": "دلار آمریکا", "EUR": "یورو"}.get(
        str(value or "").strip().upper(), str(value or "ریال"))

def _six_months_after_persian(value: str) -> str:
    match = re.fullmatch(r"(1[34]\d{2})/(\d{1,2})/(\d{1,2})", value.strip())
    if not match:
        return value
    year, month, day = map(int, match.groups())
    month += 6
    if month > 12:
        year += 1
        month -= 12
    return f"{year:04d}/{month:02d}/{day:02d}"

def _apply_contract_values(document: Document, replacements: dict) -> None:
    replacements = {**replacements, "currency": _currency_label(replacements.get("currency"))}
    start_date, end_date = str(replacements.get("startDate", "")), str(replacements.get("endDate", ""))
    first_payment_date = str(replacements.get("firstPaymentDate") or start_date).strip()
    second_payment_date = str(replacements.get("secondPaymentDate") or _six_months_after_persian(start_date)).strip()
    amount_value = str(replacements.get("amount", "")).strip()
    money_value = f"{amount_value} {replacements['currency']}".strip()
    signing_date = str(replacements.get("signingDate", "")).strip()
    signing_parts = signing_date.split("/")
    signing_date_day_first = (f"{signing_parts[2]}/{signing_parts[1]}/{signing_parts[0]}"
                              if len(signing_parts) == 3 else signing_date)
    party_name = str(replacements.get("partyName", "")).strip()
    raw_new_clauses = replacements.get("newClausesJson", "")
    try:
        new_clauses = json.loads(raw_new_clauses) if isinstance(raw_new_clauses, str) and raw_new_clauses else raw_new_clauses
    except (TypeError, ValueError):
        new_clauses = []
    if not isinstance(new_clauses, list):
        new_clauses = []
    new_clauses = [str(value).strip() for value in new_clauses if str(value).strip()]
    fallback_clause = str(replacements.get("newClause", "")).strip()
    if not new_clauses and fallback_clause:
        new_clauses = [fallback_clause]
    new_clause = "\n".join(new_clauses)
    # A single placeholder can safely hold one article. When the user asks for
    # several articles, clear it and append each article separately below.
    replacements["newClause"] = new_clause if len(new_clauses) == 1 else ""
    new_clause_number = str(replacements.get("newClauseNumber", "")).strip()
    new_clause_inserted = False
    original_contract_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
    clause_label = ("ماده" if len(re.findall(r"ماده\s*[۰-۹0-9]+", original_contract_text))
                    >= len(re.findall(r"بند\s*[۰-۹0-9]+", original_contract_text)) else "بند")
    organization_name = str(replacements.get("organizationName", "")).strip()
    organization_representative = str(replacements.get("organizationRepresentative", "")).strip()
    counterparty_name = str(replacements.get("counterpartyName", party_name)).strip()
    counterparty_representative = str(replacements.get("counterpartyRepresentative", "")).strip()
    literal_slots = {
        "فیلد1": str(replacements.get("organizationRepresentativeNationalIdentifier")
                     or replacements.get("organizationNationalIdentifier", "")),
        "فیلد2": str(replacements.get("organizationPhone", "")),
        "فیلد3": str(replacements.get("organizationAddress", "")),
        "فیلد4": "آقای" if counterparty_representative else "",
        "فیلد5": counterparty_representative,
        "فیلد6": str(replacements.get("counterpartyNationalIdentifier", "")),
        "فیلد7": str(replacements.get("counterpartyFatherName", "")),
        "فیلد8": str(replacements.get("counterpartyPhone", "")),
        "فیلد9": str(replacements.get("counterpartyAddress", "")),
        "توسعه ارتباطات": counterparty_name,
        "قیلد1": str(replacements.get("subject", "")),
        "قیلد2": start_date,
        "قیلد3": end_date,
        "قیلد4": "1 سال",
        "قیلد5": amount_value,
        "قیلد6": (f"مبلغ قرارداد در دو قسط پرداخت می‌گردد: نیمی از مبلغ در تاریخ {first_payment_date} "
                  f"و نیمی از مبلغ در تاریخ {second_payment_date}."),
        "قیلد7": str(replacements.get("visitFee", "")),
        "قیلد8": str(replacements.get("taxPercent", "")),
        "قیلد9": "14",
        "میلد1": signing_date_day_first,
        "میلد2": "",
        "میلد3": organization_representative,
        "میلد4": f"مدیرعامل شرکت {organization_name}" if organization_name else "",
        "میلد5": counterparty_representative,
        "میلد6": organization_name,
        "میلد7": organization_representative,
        "میلد8": str(replacements.get("organizationFatherName", "")),
    }
    date_pattern = re.compile(r"(?<!\d)(?:(?:13|14)\d{2}[./-]\d{1,2}[./-]\d{1,2}|\d{1,2}[./-]\d{1,2}[./-](?:13|14)\d{2})(?!\d)")
    money_pattern = re.compile(r"[۰-۹٠-٩\d]{1,3}(?:[.,٬،][۰-۹٠-٩\d]{3})+(?:\s*(?:ریال|تومان|IRR|IRT))?", re.I)

    def date_like_source(_source: str, target: str) -> str:
        # Dates across the system and generated documents use one canonical
        # Jalali representation: yyyy/MM/dd, regardless of the template sample.
        return target

    def replace_in_paragraph(paragraph):
        nonlocal new_clause_inserted
        original = ''.join(run.text or '' for run in paragraph.runs)
        combined = original
        if "{{newClause}}" in combined and new_clause:
            new_clause_inserted = True
        for key, value in replacements.items():
            combined = combined.replace("{{" + str(key) + "}}", str(value))
        had_literal_slot = False
        for key, value in literal_slots.items():
            if key in combined:
                combined = combined.replace(key, value)
                had_literal_slot = True
        if not had_literal_slot and "لغایت" in combined and start_date and end_date:
            dates = iter([start_date, end_date])
            combined = date_pattern.sub(lambda match: date_like_source(match.group(0), next(dates, match.group(0))), combined, count=2)
        if not had_literal_slot and "مبلغ قرارداد سالیانه" in combined and amount_value:
            combined = money_pattern.sub(money_value, combined, count=1)
        if not had_literal_slot and "تاریخ عقد قرارداد" in combined and start_date:
            combined = date_pattern.sub(lambda match: date_like_source(match.group(0), start_date), combined, count=1)
        if not had_literal_slot and "شش ماه بعد" in combined and second_payment_date:
            combined = date_pattern.sub(lambda match: date_like_source(match.group(0), second_payment_date), combined, count=1)
        if paragraph.runs and combined != original:
            paragraph.runs[0].text = combined
            for run in paragraph.runs[1:]:
                run.text = ''

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    def set_paragraph_text(paragraph, value: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(value)

    # Some legacy templates contain a literal subject instead of {{subject}}.
    # Treat the content immediately after "ماده 1- موضوع قرارداد" as a semantic
    # slot so a copied contract can never retain its old subject/year.
    subject = str(replacements.get("subject", "")).strip()
    subject_heading = re.compile(r"ماده\s*[۱١1]\s*[-–—:]?\s*موضوع\s+قرارداد", re.I)
    body_paragraphs = list(document.paragraphs)
    if subject:
        for index, paragraph in enumerate(body_paragraphs):
            text = paragraph.text.strip()
            heading_match = subject_heading.search(text)
            if not heading_match:
                continue
            suffix = text[heading_match.end():].strip(" \t\r\n-–—:")
            if suffix:
                set_paragraph_text(paragraph, text[:heading_match.end()] + "\n" + subject)
            else:
                next_paragraph = next((candidate for candidate in body_paragraphs[index + 1:]
                                       if candidate.text.strip()), None)
                if next_paragraph is not None:
                    set_paragraph_text(next_paragraph, subject)
            break

    # Keep the two party introductions in one continuous paragraph. Word
    # templates often store them as two paragraphs and leave a visible blank
    # line before the second party ("و فسا به نمایندگی ...").
    body_paragraphs = list(document.paragraphs)
    non_empty = [(index, paragraph) for index, paragraph in enumerate(body_paragraphs)
                 if paragraph.text.strip()]
    for pair_index in range(len(non_empty) - 1):
        first_index, first = non_empty[pair_index]
        second_index, second = non_empty[pair_index + 1]
        first_text, second_text = first.text.strip(), second.text.strip()
        if ("«کارفرما»" in first_text and first_text.endswith("و")
                and "«کارشناس»" in second_text):
            set_paragraph_text(first, first_text + " " + second_text)
            for paragraph in body_paragraphs[first_index + 1:second_index + 1]:
                paragraph._element.getparent().remove(paragraph._element)
            break
    if new_clauses and not new_clause_inserted:
        article_pattern = re.compile(r"^\s*" + re.escape(clause_label) + r"\s*[۰-۹0-9]+")
        article_paragraphs = [p for p in document.paragraphs if article_pattern.search(p.text)]
        for table in document.tables:
            article_paragraphs.extend(
                p for row in table.rows for cell in row.cells for p in cell.paragraphs
                if article_pattern.search(p.text))
        terminal_markers = ("تعداد نسخ", "امضای طرفین", "امضاء طرفین", "تاریخ")
        terminal_article = next((p for p in reversed(article_paragraphs)
            if sum(marker in p.text for marker in terminal_markers) >= 2), None)

        effective_number = int(new_clause_number) if new_clause_number.isdigit() else None
        if terminal_article is not None:
            number_match = re.search(r"[۰-۹0-9]+", terminal_article.text)
            if number_match:
                latin_number = number_match.group(0).translate(str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789"))
                terminal_number = int(latin_number)
                effective_number = terminal_number
                shifted_number = str(terminal_number + len(new_clauses))

                def replace_paragraph_text(target, pattern, replacement):
                    original = "".join(run.text or "" for run in target.runs)
                    updated = re.sub(pattern, replacement, original, count=1)
                    if target.runs and updated != original:
                        target.runs[0].text = updated
                        for extra_run in target.runs[1:]:
                            extra_run.text = ""

                replace_paragraph_text(terminal_article, r"[۰-۹0-9]+", shifted_number)
                count_pattern = rf"(در\s*){terminal_number}(\s*{re.escape(clause_label)})"
                for target in document.paragraphs:
                    replace_paragraph_text(target, count_pattern, rf"\g<1>{shifted_number}\g<2>")
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for target in cell.paragraphs:
                                replace_paragraph_text(target, count_pattern,
                                                       rf"\g<1>{shifted_number}\g<2>")

        substantive_articles = [p for p in article_paragraphs if p is not terminal_article]
        style_source = (substantive_articles[-1] if substantive_articles
                        else article_paragraphs[-1] if article_paragraphs else None)
        signature_markers = ("امضا", "امضاء", "مدیرعامل", "نماینده طرف", "طرف اول", "طرف دوم")
        body_elements = list(document.element.body.iterchildren())
        insertion_target = None
        for element in body_elements:
            block_text = "".join(element.itertext())
            if any(marker in block_text for marker in signature_markers):
                insertion_target = element
        for index, clause_text in enumerate(new_clauses):
            clause_number = effective_number + index if effective_number is not None else None
            label = (f"{clause_label} {clause_number} - " if clause_number is not None
                     else f"{clause_label} جدید - ")
            paragraph = document.add_paragraph()
            if style_source is not None:
                if style_source._p.pPr is not None:
                    paragraph._p.insert(0, deepcopy(style_source._p.pPr))
                run = paragraph.add_run(f"{label}{clause_text}")
                if style_source.runs and style_source.runs[0]._r.rPr is not None:
                    run._r.insert(0, deepcopy(style_source.runs[0]._r.rPr))
            else:
                paragraph.add_run(f"{label}{clause_text}")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if terminal_article is not None:
                terminal_article._p.addprevious(paragraph._p)
            elif insertion_target is not None:
                insertion_target.addprevious(paragraph._p)

@app.post("/contract/generate")
async def generate_contract(file: UploadFile = File(...), values: str = Form("{}")):
    import json
    try:
        replacements = json.loads(values)
        source = BytesIO(await file.read())
        document = Document(source)
        _apply_contract_values(document, replacements)
        output = BytesIO()
        document.save(output)
        return Response(output.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=generated-contract.docx"})
    except Exception as exc:
        raise HTTPException(422, f"Contract generation failed: {exc}") from exc

def _contract_pdf(payload: dict) -> bytes:
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from xml.sax.saxutils import escape

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    pdfmetrics.registerFont(TTFont("ContractDejaVu", font_path))
    def fa(value):
        return escape(get_display(reshape(_report_text(value))))
    output = BytesIO()
    document = SimpleDocTemplate(output, pagesize=A4, rightMargin=18*mm,
        leftMargin=18*mm, topMargin=17*mm, bottomMargin=18*mm,
        title="NegareshAI Contract Draft")
    base = getSampleStyleSheet()
    body = ParagraphStyle("ContractBody", parent=base["BodyText"],
        fontName="ContractDejaVu", fontSize=9, leading=16,
        alignment=TA_RIGHT, spaceAfter=6)
    title = ParagraphStyle("ContractTitle", parent=body, fontSize=17,
        leading=25, textColor=colors.HexColor("#0B2545"), spaceAfter=10)
    heading = ParagraphStyle("ContractHeading", parent=body, fontSize=12,
        leading=19, textColor=colors.HexColor("#2E5AAC"), spaceBefore=9,
        spaceAfter=6)
    story = [Paragraph(fa("پیش‌نویس هوشمند قرارداد"), title),
        Paragraph(fa(payload.get("subject")), heading)]
    metadata = [
        ("طرف قرارداد", payload.get("partyName")),
        ("تاریخ شروع", payload.get("startDate")),
        ("تاریخ پایان", payload.get("endDate")),
        ("مبلغ", f"{_report_text(payload.get('amount'))} {_currency_label(payload.get('currency'))}"),
        ("نسخه پیش‌نویس", payload.get("draftVersion")),
        ("شناسه گفت‌وگو", payload.get("conversationId")),
    ]
    rows = [[Paragraph(fa(label), body), Paragraph(fa(value), body)] for label, value in metadata]
    table = Table(rows, colWidths=[42*mm, 126*mm])
    table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), .35, colors.HexColor("#D9DEE7")),
        ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#F2F4F7")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("RIGHTPADDING", (0,0), (-1,-1), 7),
        ("TOPPADDING", (0,0), (-1,-1), 5),
        ("BOTTOMPADDING", (0,0), (-1,-1), 5),
    ]))
    story.extend([table, Spacer(1, 8)])
    approved = _report_text(payload.get("approvedClauses"))
    direct = _report_text(payload.get("newClause"))
    if approved:
        story.extend([Paragraph(fa("بندهای مصوب گروه قرارداد"), heading),
            Paragraph(fa(approved).replace("\n", "<br/>"), body)])
    if direct:
        story.extend([Paragraph(fa("بند ناشی از دستور مستقیم کاربر"), heading),
            Paragraph(fa(direct).replace("\n", "<br/>"), body)])
    story.append(Paragraph(fa("منابع و استنادها"), heading))
    citations = payload.get("citations") or []
    if not citations:
        story.append(Paragraph(fa("این پیش‌نویس منبع بازیابی‌شده‌ای ندارد و بر مبنای قالب و پاسخ‌های کاربر تولید شده است."), body))
    for index, citation in enumerate(citations, start=1):
        label = (f"{index}. {_report_text(citation.get('documentTitle'))}، "
                 f"صفحه {_report_text(citation.get('page'))}، "
                 f"بخش {_report_text(citation.get('section'))}")
        story.append(Paragraph(fa(label), body))
        story.append(Paragraph(fa(citation.get("evidence")).replace("\n", "<br/>"), body))
    story.extend([Paragraph(fa("خلاصه تغییرات ساختاریافته"), heading),
        Paragraph(fa(payload.get("diffJson")), body)])
    def footer(canvas, doc):
        canvas.saveState(); canvas.setFont("ContractDejaVu", 7)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawCentredString(A4[0]/2, 9*mm,
            get_display(reshape(f"نسخه {payload.get('draftVersion')} | صفحه {doc.page}")))
        canvas.restoreState()
    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()

@app.post("/contract/pdf")
async def generate_contract_pdf(payload: dict):
    try:
        return Response(_contract_pdf(payload), media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=contract-draft.pdf"})
    except Exception as exc:
        raise HTTPException(422, f"Contract PDF generation failed: {exc}") from exc

@app.post("/contract/compare")
async def compare_contracts(original: UploadFile = File(...), revised: UploadFile = File(...)):
    async def read_doc(upload: UploadFile) -> list[str]:
        document = Document(BytesIO(await upload.read()))
        return [p.text for p in document.paragraphs]
    try:
        before, after = await read_doc(original), await read_doc(revised)
        changes = []
        for item in difflib.ndiff(before, after):
            if item.startswith(('+ ', '- ')):
                changes.append({"type": "added" if item.startswith('+ ') else "removed", "text": item[2:]})
        return {"changes": changes, "added": sum(x["type"] == "added" for x in changes), "removed": sum(x["type"] == "removed" for x in changes)}
    except Exception as exc:
        raise HTTPException(422, f"Contract comparison failed: {exc}") from exc

@app.post("/compliance/check")
async def compliance_check(payload: dict):
    text = str(payload.get("text", ""))
    checklist = payload.get("checklist", [])
    if not text.strip() or not isinstance(checklist, list) or not checklist:
        raise HTTPException(400, "text and checklist are required")
    findings = []
    pages = text.split("\f")
    def locate(value: str):
        offset = 0
        position = text.find(value)
        if position < 0:
            return None, None
        for page_number, page_text in enumerate(pages, start=1):
            if position <= offset + len(page_text):
                local = max(0, position - offset)
                start = max(0, local - 120)
                return page_text[start:local + len(value) + 120], page_number
            offset += len(page_text) + 1
        return value, 1
    for item in checklist:
        requirement = str(item.get("requirement", item) if isinstance(item, dict) else item).strip()
        if not requirement:
            continue
        forbidden = bool(item.get("forbidden", False)) if isinstance(item, dict) else False
        weight = max(0.0, float(item.get("weight", 1))) if isinstance(item, dict) else 1.0
        critical = bool(item.get("critical", False)) if isinstance(item, dict) else False
        evidence, page = locate(requirement)
        present = evidence is not None
        passed = not present if forbidden else present
        findings.append({"code": item.get("code") if isinstance(item, dict) else None,
            "requirement": requirement, "status": "met" if passed else ("forbidden" if forbidden else "missing"),
            "passed": passed, "weight": weight, "critical": critical,
            "evidence": evidence, "page": page, "section": item.get("section") if isinstance(item, dict) else None,
            "confidence": 1.0 if present else 0.92,
            "suggestion": None if passed else (f"عبارت ممنوع «{requirement}» حذف شود." if forbidden else f"الزام «{requirement}» افزوده شود.")})
    applicable_weight = sum(x["weight"] for x in findings)
    passed_weight = sum(x["weight"] for x in findings if x["passed"])
    weighted_score = round((passed_weight * 100 / applicable_weight), 2) if applicable_weight else 0
    critical_failure = any(x["critical"] and not x["passed"] for x in findings)
    threshold = min(100.0, max(0.0, float(payload.get("passingThreshold", 80))))
    decision = "non_compliant" if critical_failure or weighted_score < threshold else "compliant"
    # Reflection pass: every emitted citation must be recoverable from the target text.
    citation_failures = [x for x in findings if x["evidence"] and x["evidence"] not in text]
    if citation_failures:
        decision = "needs_human_review"
        for finding in citation_failures:
            finding["confidence"] = min(finding["confidence"], 0.49)
    missing = [x for x in findings if not x["passed"]]
    focus = [str(x) for x in payload.get("focus", []) if str(x).strip()]
    focus_findings = [{"topic": topic, "present": topic in text, "evidence": text[max(0, text.find(topic)-120):text.find(topic)+len(topic)+120] if topic in text else None} for topic in focus]
    return {"decision": decision, "weightedScore": weighted_score,
        "passingThreshold": threshold, "criticalFailure": critical_failure,
        "findings": findings, "focusFindings": focus_findings,
        "missingCount": len(missing), "totalCount": len(findings),
        "toolTrace": {"strategy": "two-pass-evidence-grounded",
            "tools": ["criterion-evaluator", "page-citation-locator", "weighted-score-calculator", "citation-reflection-verifier"],
            "reflection": {"passes": 2, "citationFailures": len(citation_failures)},
            "mcp": {"used": False, "reason": "private supplied text and deterministic tools were sufficient"}}}

def _report_text(value) -> str:
    return "" if value is None else str(value)

def _comparison_docx(payload: dict) -> bytes:
    from docx import Document as WordDocument
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = WordDocument()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name, normal.font.size = "Arial", Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        style = styles[name]
        style.font.name, style.font.size = "Arial", Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run("NegareshAI | گزارش محرمانه تطابق اسناد")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"شناسه اجرا: {_report_text(payload.get('id'))}")

    def rtl(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        props = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        props.append(bidi)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
        return paragraph

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("گزارش ممیزی تطابق اسناد")
    run.bold, run.font.size = True, Pt(23)
    rtl(title)
    subtitle = document.add_paragraph(
        f"{_report_text(payload.get('targetDocumentTitle'))} | "
        f"{_report_text(payload.get('createdAtLabel'))}")
    subtitle.paragraph_format.space_after = Pt(16)
    rtl(subtitle)

    metadata = [
        ("نتیجه کلی", payload.get("outcomeLabel")),
        ("امتیاز تطابق", f"{_report_text(payload.get('scorePercent'))}%"),
        ("آستانه قبولی", f"{_report_text(payload.get('passingThreshold'))}%"),
        ("نقض معیار حیاتی", "بله" if payload.get("hasCriticalFailure") else "خیر"),
        ("توضیح نتیجه", payload.get("outcomeExplanation")),
        ("مبنای تطابق", payload.get("basisLabel")),
        ("مدل محلی", payload.get("modelId")),
        ("نسخه prompt", payload.get("promptVersion")),
        ("نسخه هدف", payload.get("targetVersionId")),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(_report_text(value))
        rtl(paragraph)

    heading = document.add_heading("یافته‌ها و شواهد", level=1)
    rtl(heading)
    findings = payload.get("findings") or []
    if not findings:
        rtl(document.add_paragraph("یافته‌ای برای این اجرا ثبت نشده است."))
    for index, finding in enumerate(findings, start=1):
        heading = document.add_heading(
            f"{index}. {_report_text(finding.get('title'))}", level=2)
        rtl(heading)
        table = document.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width, table.columns[1].width = Inches(1.45), Inches(5.05)
        rows = [
            ("وضعیت", finding.get("typeLabel")),
            ("شدت", finding.get("severity")),
            ("وزن", finding.get("weight")),
            ("معیار حیاتی", "بله" if finding.get("isCritical") else "خیر"),
            ("دلیل", finding.get("reason")),
            ("شاهد هدف", finding.get("targetEvidence")),
            ("صفحه هدف", finding.get("targetPage")),
            ("شاهد مرجع", finding.get("referenceEvidence")),
            ("صفحه مرجع", finding.get("referencePage")),
            ("بخش مرجع", finding.get("referenceSection")),
            ("پیشنهاد اصلاح", finding.get("suggestion")),
            ("اطمینان", finding.get("confidence")),
            ("تصمیم کارشناس", finding.get("reviewLabel")),
            ("نظر کارشناس", finding.get("reviewerComment")),
        ]
        for label, value in rows:
            row = table.add_row()
            row.cells[0].width, row.cells[1].width = Inches(1.45), Inches(5.05)
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].text, row.cells[1].text = label, _report_text(value)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    rtl(paragraph)
        document.add_paragraph()
    output = BytesIO()
    document.save(output)
    return output.getvalue()

def _comparison_pdf(payload: dict) -> bytes:
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    pdfmetrics.registerFont(TTFont("DejaVu", font_path))
    def fa(value):
        return get_display(reshape(_report_text(value)))
    output = BytesIO()
    document = SimpleDocTemplate(output, pagesize=A4, rightMargin=20*mm,
        leftMargin=20*mm, topMargin=18*mm, bottomMargin=18*mm,
        title="NegareshAI Comparison Audit Report")
    base = getSampleStyleSheet()
    body = ParagraphStyle("PersianBody", parent=base["BodyText"], fontName="DejaVu",
        fontSize=9, leading=15, alignment=TA_RIGHT, spaceAfter=5)
    title = ParagraphStyle("PersianTitle", parent=body, fontSize=18,
        leading=25, textColor=colors.HexColor("#0B2545"), spaceAfter=8)
    heading = ParagraphStyle("PersianHeading", parent=body, fontSize=12,
        leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=10,
        spaceAfter=6)
    story = [
        Paragraph(fa("گزارش ممیزی تطابق اسناد"), title),
        Paragraph(fa(payload.get("targetDocumentTitle")), heading),
        Paragraph(fa(f"تاریخ اجرا: {payload.get('createdAtLabel')}"), body),
        Paragraph(fa(f"شناسه اجرا: {payload.get('id')}"), body),
        Paragraph(fa(f"نتیجه: {payload.get('outcomeLabel')} | امتیاز: {payload.get('scorePercent')} درصد | آستانه: {payload.get('passingThreshold')} درصد"), body),
        Paragraph(fa(f"نقض معیار حیاتی: {'بله' if payload.get('hasCriticalFailure') else 'خیر'} | {payload.get('outcomeExplanation')}"), body),
        Paragraph(fa(f"مدل: {payload.get('modelId')} | نسخه prompt: {payload.get('promptVersion')}"), body),
        Spacer(1, 8),
    ]
    for index, finding in enumerate(payload.get("findings") or [], start=1):
        story.append(Paragraph(fa(f"{index}. {finding.get('title')}"), heading))
        data = [[Paragraph(fa(label), body), Paragraph(fa(value), body)] for label, value in [
            ("وضعیت", finding.get("typeLabel")), ("شدت", finding.get("severity")),
            ("وزن", finding.get("weight")),
            ("معیار حیاتی", "بله" if finding.get("isCritical") else "خیر"),
            ("دلیل", finding.get("reason")),
            ("شاهد هدف", finding.get("targetEvidence")),
            ("صفحه هدف", finding.get("targetPage")),
            ("شاهد مرجع", finding.get("referenceEvidence")),
            ("پیشنهاد اصلاح", finding.get("suggestion")),
            ("اطمینان", finding.get("confidence")),
            ("تصمیم کارشناس", finding.get("reviewLabel")),
        ]]
        table = Table(data, colWidths=[35*mm, 125*mm], repeatRows=0)
        table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), .35, colors.HexColor("#D9DEE7")),
            ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#F2F4F7")),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("RIGHTPADDING", (0,0), (-1,-1), 7),
            ("LEFTPADDING", (0,0), (-1,-1), 7),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ]))
        story.extend([table, Spacer(1, 8)])
    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("DejaVu", 7)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawCentredString(
            A4[0] / 2, 9*mm,
            fa(f"شناسه اجرا: {payload.get('id')} | صفحه {doc.page}"))
        canvas.restoreState()
    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()

@app.post("/comparison/report")
async def comparison_report(payload: dict, format: str = "docx"):
    if format not in {"docx", "pdf"}:
        raise HTTPException(400, "format must be docx or pdf")
    try:
        content = _comparison_docx(payload) if format == "docx" else _comparison_pdf(payload)
        media_type = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                      if format == "docx" else "application/pdf")
        return Response(content, media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename=comparison-report.{format}"})
    except Exception as exc:
        raise HTTPException(422, f"Comparison report generation failed: {exc}") from exc
