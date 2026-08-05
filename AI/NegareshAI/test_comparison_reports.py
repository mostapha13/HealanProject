import os
import unittest
import asyncio
from io import BytesIO
from pathlib import Path

from docx import Document

import main


class ComparisonReportTests(unittest.TestCase):
    def setUp(self):
        self.payload = {
            "id": "11111111-2222-3333-4444-555555555555",
            "targetDocumentTitle": "امیدنامه شرکت فولاد دهدشت",
            "targetVersionId": "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee",
            "createdAtUtc": "2026-07-30T08:30:00Z",
            "outcomeLabel": "نیازمند بررسی انسانی",
            "scorePercent": 66.67,
            "basisLabel": "ترکیبی",
            "modelId": "BAAI/bge-m3",
            "promptVersion": "comparison.prompt:v1",
            "findings": [
                {
                    "title": "بند محرمانگی",
                    "typeLabel": "مفقود",
                    "severity": 4,
                    "reason": "بند محرمانگی در سند هدف یافت نشد.",
                    "targetEvidence": "موضوع فعالیت شرکت تولید محصولات فولادی است.",
                    "targetPage": 2,
                    "referenceEvidence": "اطلاعات طرفین محرمانه تلقی می‌شود.",
                    "referencePage": 5,
                    "suggestion": "بند محرمانگی مطابق سند مرجع اضافه شود.",
                    "confidence": 0.95,
                    "reviewLabel": "در انتظار بررسی",
                    "reviewerComment": None,
                },
                {
                    "title": "سرمایه ثبت‌شده",
                    "typeLabel": "منطبق",
                    "severity": 2,
                    "reason": "مبلغ سرمایه با قاعده ثبت‌شده منطبق است.",
                    "targetEvidence": "سرمایه ثبت‌شده ۱۵٬۰۰۰٬۰۰۰٬۰۰۰ ریال است.",
                    "targetPage": 3,
                    "referenceEvidence": None,
                    "referencePage": None,
                    "suggestion": None,
                    "confidence": 0.98,
                    "reviewLabel": "تأییدشده",
                    "reviewerComment": "مبلغ با صورت مالی تطبیق داده شد.",
                },
            ],
        }

    def test_docx_is_readable_and_contains_audit_content(self):
        content = main._comparison_docx(self.payload)
        self.assertTrue(content.startswith(b"PK"))
        document = Document(BytesIO(content))
        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables
               for row in table.rows for cell in row.cells]
        )
        self.assertIn("گزارش ممیزی تطابق اسناد", all_text)
        self.assertIn("بند محرمانگی", all_text)
        self._write_artifact("p3-comparison-report.docx", content)

    def test_pdf_has_valid_header_and_persian_content(self):
        content = main._comparison_pdf(self.payload)
        self.assertTrue(content.startswith(b"%PDF-"))
        self.assertGreater(len(content), 5_000)
        self._write_artifact("p3-comparison-report.pdf", content)

    def test_weighted_critical_check_uses_reflection_and_page_citations(self):
        result = asyncio.run(main.compliance_check({
            "text": "نام شرکت فولاد دهدشت\fسرمایه ثبت‌شده شرکت",
            "passingThreshold": 80,
            "checklist": [
                {"code": "NAME", "requirement": "نام شرکت", "weight": 90},
                {"code": "SECRET", "requirement": "بند محرمانگی", "weight": 10, "critical": True},
            ],
        }))
        self.assertEqual(90, result["weightedScore"])
        self.assertTrue(result["criticalFailure"])
        self.assertEqual("non_compliant", result["decision"])
        self.assertEqual(1, result["findings"][0]["page"])
        self.assertEqual(2, result["toolTrace"]["reflection"]["passes"])

    @staticmethod
    def _write_artifact(name: str, content: bytes):
        output = os.getenv("REPORT_TEST_OUTPUT")
        if not output:
            return
        path = Path(output)
        path.mkdir(parents=True, exist_ok=True)
        (path / name).write_bytes(content)


if __name__ == "__main__":
    unittest.main()
