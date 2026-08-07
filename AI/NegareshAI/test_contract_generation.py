import unittest
import json
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from main import app


class ContractGenerationTests(unittest.TestCase):
    def test_appends_numbered_clause_when_template_has_no_placeholder(self):
        template = Document()
        template.add_paragraph("بند 10 - آخرین بند قرارداد")
        source = BytesIO()
        template.save(source)

        response = self.client.post("/contract/generate", files={
            "file": ("template.docx", source.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": json.dumps({
            "newClause": "پیگیری شکایت از دیوان عدالت اداری",
            "newClauseNumber": "11"
        }, ensure_ascii=False)})

        self.assertEqual(200, response.status_code, response.text)
        generated = Document(BytesIO(response.content))
        self.assertIn("بند 11 - پیگیری شکایت از دیوان عدالت اداری",
                      "\n".join(x.text for x in generated.paragraphs))

    def test_inserts_new_article_before_signature_block(self):
        template = Document()
        last_article = template.add_paragraph("ماده 13 - حل اختلاف")
        last_article.style = template.styles["Heading 2"]
        last_article.runs[0].bold = True
        template.add_paragraph("متن و جزئیات ماده سیزدهم")
        template.add_paragraph("ماده 14 - تعداد نسخ / امضای طرفین / تاریخ")
        template.add_paragraph("این قرارداد در 14 ماده تنظیم و امضا گردید.")
        signature = template.add_table(rows=1, cols=2)
        signature.cell(0, 0).text = "امضاء طرف اول"
        signature.cell(0, 1).text = "مدیرعامل شرکت فسا"
        source = BytesIO()
        template.save(source)

        response = self.client.post("/contract/generate", files={
            "file": ("template.docx", source.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": json.dumps({
            "newClause": "کلیه دعاوی و اختلافات از طریق دعوای فیزیکی قابل دریافت است",
            "newClauseNumber": "15"
        }, ensure_ascii=False)})

        self.assertEqual(200, response.status_code, response.text)
        generated = Document(BytesIO(response.content))
        body = list(generated.element.body.iterchildren())
        clause = next(x for x in body if "ماده 14 - کلیه دعاوی" in "".join(x.itertext()))
        terminal = next(x for x in body if "ماده 15 - تعداد نسخ" in "".join(x.itertext()))
        signatures = next(x for x in body if "امضاء طرف اول" in "".join(x.itertext()))
        self.assertLess(body.index(clause), body.index(terminal))
        self.assertLess(body.index(clause), body.index(signatures))
        self.assertIn("در 15 ماده", "\n".join(p.text for p in generated.paragraphs))
        clause_paragraph = next(p for p in generated.paragraphs if "ماده 14 - کلیه دعاوی" in p.text)
        self.assertEqual(last_article.style.name, clause_paragraph.style.name)
        self.assertTrue(clause_paragraph.runs[0].bold)

    def setUp(self):
        self.client = TestClient(app)

    def test_docx_template_replacement_and_source_backed_pdf(self):
        template = Document()
        template.add_paragraph("موضوع: {{subject}}")
        template.add_paragraph("مبلغ: {{amount}}")
        content = BytesIO()
        template.save(content)
        docx = self.client.post("/contract/generate", files={
            "file": ("template.docx", content.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": '{"subject":"پشتیبانی فسا","amount":"130,000,000"}'})
        self.assertEqual(200, docx.status_code)
        generated = Document(BytesIO(docx.content))
        self.assertIn("پشتیبانی فسا", "\n".join(x.text for x in generated.paragraphs))

        pdf = self.client.post("/contract/pdf", json={
            "conversationId": "11111111-1111-1111-1111-111111111111",
            "draftVersion": 2,
            "subject": "پشتیبانی شرکت فسا",
            "partyName": "شرکت فسا",
            "startDate": "1405/01/01",
            "endDate": "1405/12/29",
            "amount": "130,000,000",
            "currency": "IRR",
            "newClause": "حل اختلاف از طریق شورای سازمان",
            "approvedClauses": "بند محرمانگی مصوب",
            "diffJson": "{}",
            "createdAtUtc": "2026-08-01T00:00:00Z",
            "citations": [{
                "documentTitle": "قرارداد نهایی 1404",
                "documentId": "22222222-2222-2222-2222-222222222222",
                "versionId": "33333333-3333-3333-3333-333333333333",
                "page": 4,
                "section": "مبلغ قرارداد",
                "evidence": "مبلغ قرارداد مرجع"
            }]
        })
        self.assertEqual(200, pdf.status_code)
        self.assertTrue(pdf.content.startswith(b"%PDF"))
        self.assertGreaterEqual(len(PdfReader(BytesIO(pdf.content)).pages), 1)

    def test_real_contract_without_placeholders_updates_dates_amount_and_currency(self):
        template = Document()
        template.add_paragraph("از تاریخ 01/01/1404 لغایت 29/12/1404 به مدت 1 سال شمسی می‌باشد.")
        template.add_paragraph("مبلغ قرارداد سالیانه 19.000.000.000 ریال می‌باشد.")
        template.add_paragraph("01/01/1404 نیمی از مبلغ قرارداد در تاریخ عقد قرارداد")
        template.add_paragraph("نیمی از مبلغ قرارداد شش ماه بعد در تاریخ 01/06/1404")
        content = BytesIO()
        template.save(content)

        response = self.client.post("/contract/generate", files={
            "file": ("template.docx", content.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": '{"startDate":"1405/01/01","endDate":"1405/12/29",'
                                '"amount":"18,750,000,000","currency":"IRR"}'})

        self.assertEqual(200, response.status_code)
        generated = Document(BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
        self.assertIn("1405/01/01", text)
        self.assertIn("1405/12/29", text)
        self.assertIn("1405/07/01", text)
        self.assertIn("18,750,000,000 ریال", text)
        self.assertNotIn("19.000.000.000", text)
        self.assertNotIn("IRR", text)

    def test_fasa_literal_template_fields_are_all_filled(self):
        template = Document()
        template.add_paragraph("شرکت میلد6 نماینده میلد7 پدر میلد8 شماره ملی فیلد1 تلفن فیلد2 آدرس فیلد3")
        template.add_paragraph("توسعه ارتباطات نماینده فیلد4 فیلد5 ملی فیلد6 پدر فیلد7 تلفن فیلد8 آدرس فیلد9")
        template.add_paragraph("قیلد1")
        template.add_paragraph("از تاریخ قیلد2 لغایت قیلد3 به مدت قیلد4 شمسی")
        template.add_paragraph("مبلغ قرارداد سالیانه قیلد5 ریال")
        template.add_paragraph("قیلد6")
        template.add_paragraph("مراجعه قیلد7 ریال و مالیات قیلد8 درصد")
        template.add_paragraph("در قیلد9 ماده در تاریخ میلد1 در شهر میلد2")
        template.add_paragraph("میلد3 میلد4 میلد5")
        content = BytesIO()
        template.save(content)
        response = self.client.post("/contract/generate", files={
            "file": ("Template.docx", content.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": '{"subject":"قرارداد پشتیبانی فسا","partyName":"شرکت فسا",'
                                '"organizationName":"داده پردازان","organizationRepresentative":"مصطفی مهدوی",'
                                '"organizationFatherName":"ابراهیم","organizationRepresentativeNationalIdentifier":"0012345678",'
                                '"organizationNationalIdentifier":"14001234567","counterpartyName":"فسا",'
                                '"counterpartyRepresentative":"محمد محمدی","counterpartyNationalIdentifier":"11125456",'
                                '"counterpartyPhone":"09122222221",'
                                '"startDate":"1402/01/01","endDate":"1402/12/29",'
                                '"amount":"6,000,000,000","currency":"ریال",'
                                '"signingDate":"1405/05/15"}'})
        self.assertEqual(200, response.status_code)
        generated = Document(BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
        self.assertNotRegex(text, r"(?:فیلد|قیلد)[1-9]|میلد[1-5]")
        self.assertIn("1402/01/01", text)
        self.assertIn("1402/12/29", text)
        self.assertIn("6,000,000,000 ریال", text)
        self.assertIn("15/05/1405", text)
        self.assertIn("شرکت داده پردازان نماینده مصطفی مهدوی پدر ابراهیم", text)
        self.assertIn("شماره ملی 0012345678", text)
        self.assertNotIn("شماره ملی 14001234567", text)
        self.assertIn("فسا نماینده آقای محمد محمدی ملی 11125456", text)
        self.assertNotIn("سارا سارایی", text)

    def test_replaces_literal_subject_and_joins_party_intro_paragraphs(self):
        template = Document()
        template.add_paragraph(
            "شرکت داده پردازان هوشمند مهر که ازین پس «کارفرما» نامیده خواهد شد و")
        template.add_paragraph("")
        template.add_paragraph(
            "فسا به نمایندگی آقای محمد محمدی که ازین پس «کارشناس» نامیده خواهد شد")
        template.add_paragraph("ماده 1- موضوع قرارداد")
        template.add_paragraph("قرارداد پشتیبانی شرکت فسا 1403")
        content = BytesIO()
        template.save(content)

        response = self.client.post("/contract/generate", files={
            "file": ("Template.docx", content.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": json.dumps({
            "subject": "پشتیبانی شرکت فسا"
        }, ensure_ascii=False)})

        self.assertEqual(200, response.status_code, response.text)
        generated = Document(BytesIO(response.content))
        paragraphs = [paragraph.text for paragraph in generated.paragraphs]
        text = "\n".join(paragraphs)
        self.assertIn("پشتیبانی شرکت فسا", text)
        self.assertNotIn("قرارداد پشتیبانی شرکت فسا 1403", text)
        intro = next(value for value in paragraphs if "«کارفرما»" in value)
        self.assertIn("و فسا به نمایندگی", intro)
        self.assertEqual(1, sum("«کارشناس»" in value for value in paragraphs))

    def test_applies_two_articles_and_explicit_payment_dates(self):
        template = Document()
        template.add_paragraph("قیلد6")
        template.add_paragraph("ماده 13- حل اختلاف")
        template.add_paragraph("متن حل اختلاف")
        template.add_paragraph("ماده 14- تعداد نسخ / امضای طرفین/ تاریخ")
        template.add_paragraph("این قرارداد در 14 ماده تنظیم شد")
        content = BytesIO()
        template.save(content)

        clauses = ["کلیه دعاوی از طریق مرجع تعیین‌شده رسیدگی می‌شود",
                   "کلیه اسناد باید محرمانه باشد"]
        response = self.client.post("/contract/generate", files={
            "file": ("Template.docx", content.getvalue(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }, data={"values": json.dumps({
            "startDate": "1405/01/01", "endDate": "1405/12/29",
            "firstPaymentDate": "1405/01/01", "secondPaymentDate": "1405/06/01",
            "newClauseNumber": "14", "newClause": clauses[-1],
            "newClausesJson": json.dumps(clauses, ensure_ascii=False)
        }, ensure_ascii=False)})

        self.assertEqual(200, response.status_code)
        generated = Document(BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
        self.assertIn("ماده 14 - " + clauses[0], text)
        self.assertIn("ماده 15 - " + clauses[1], text)
        self.assertIn("ماده 16- تعداد نسخ", text)
        self.assertIn("در 16 ماده", text)
        self.assertIn("تاریخ 1405/01/01", text)
        self.assertIn("تاریخ 1405/06/01", text)


if __name__ == "__main__":
    unittest.main()
